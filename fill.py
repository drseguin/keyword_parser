import streamlit as st
import os
import re
import docx
import tempfile
from excel_manager import excelManager
from keyword_parser import keywordParser

def process_word_doc(doc_path, excel_path):
    """
    Process a Word document, replacing keywords with values from Excel spreadsheet.
    
    Args:
        doc_path: Path to the Word document
        excel_path: Path to the Excel spreadsheet
        
    Returns:
        Processed document object and a count of replaced keywords
    """
    # Load the document
    doc = docx.Document(doc_path)
    
    # Initialize Excel manager
    excel_mgr = excelManager(excel_path)
    
    # Initialize keyword parser with the Excel manager and pass the document reference
    parser = keywordParser(excel_mgr)
    parser.set_word_document(doc)
    
    # Compile regex pattern for keywords
    pattern = r'{{(.*?)}}'
    
    # Count total keywords for progress tracking
    total_keywords = 0
    
    # Count keywords in paragraphs
    for paragraph in doc.paragraphs:
        total_keywords += len(re.findall(pattern, paragraph.text))
    
    # Count keywords in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    total_keywords += len(re.findall(pattern, paragraph.text))
    
    if total_keywords == 0:
        st.warning("No keywords found in the document.")
        return doc, 0
    
    # Initialize progress bar
    progress_bar = st.progress(0)
    progress_text = st.empty()
    
    # Counter for processed keywords
    processed_count = 0
    
    # First handle all INPUT keywords with a single form if needed
    input_keywords = []
    input_locations = []
    
    # Collect all INPUT keywords from paragraphs
    for paragraph in doc.paragraphs:
        matches = list(re.finditer(pattern, paragraph.text))
        for match in matches:
            keyword = match.group(0)
            content = match.group(1)
            if content.split(":", 1)[0].strip().upper() == "INPUT":
                input_keywords.append(keyword)
                input_locations.append(("paragraph", paragraph, match.start(), match.end()))
    
    # Collect all INPUT keywords from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    matches = list(re.finditer(pattern, paragraph.text))
                    for match in matches:
                        keyword = match.group(0)
                        content = match.group(1)
                        if content.split(":", 1)[0].strip().upper() == "INPUT":
                            input_keywords.append(keyword)
                            input_locations.append(("table_cell", paragraph, match.start(), match.end()))
    
    # If there are input keywords, process them all at once
    input_values = {}
    if input_keywords:
        with st.form(key="document_input_form"):
            st.subheader("Please provide values for input fields:")
            
            for keyword in input_keywords:
                # Extract the field name for better labels
                content = re.match(pattern, keyword).group(1)
                parts = content.split(":", 1)
                field_name = parts[0].strip().replace("INPUT:", "")
                
                # Create appropriate input fields
                if len(parts) > 1 and "select:" in parts[1]:
                    options = parts[1].split("select:")[1].split(",")
                    options = [opt.strip() for opt in options]
                    value = st.selectbox(f"Select for {field_name}", options)
                elif len(parts) > 1 and "date:" in parts[1]:
                    value = st.date_input(f"Date for {field_name}")
                    value = value.strftime("%Y-%m-%d")
                else:
                    default = parts[1].strip() if len(parts) > 1 else ""
                    value = st.text_input(f"Enter {field_name}", value=default)
                
                input_values[keyword] = value
            
            submit = st.form_submit_button("Submit")
            if not submit:
                st.stop()  # Stop execution until form is submitted
    
    # List to track table placeholders
    table_placeholders = []
    table_placeholder_index = 0
    
    # Track paragraphs with Excel range keywords to handle special table insertion
    range_paragraphs = []
    
    # First scan to find paragraphs containing Excel range keywords
    for paragraph in doc.paragraphs:
        matches = list(re.finditer(pattern, paragraph.text))
        for match in matches:
            content = match.group(1)
            parts = content.split(":", 1)
            if parts[0].strip().upper() == "XL" and len(parts) > 1:
                # Check if it's a range reference like A1:C3
                if ":" in parts[1] and "!" not in parts[1].split(":", 1)[1]:
                    # Store the paragraph for special handling
                    if paragraph not in range_paragraphs:
                        range_paragraphs.append(paragraph)
    
    # Now process all keywords in paragraphs, with special handling for range paragraphs
    for paragraph in doc.paragraphs:
        # If this is a range paragraph that needs special handling for table insertion
        if paragraph in range_paragraphs:
            # Get all keywords in this paragraph
            matches = list(re.finditer(pattern, paragraph.text))
            
            for match in matches:
                keyword = match.group(0)
                content = match.group(1)
                parts = content.split(":", 1)
                
                # Process Excel range keywords
                if parts[0].strip().upper() == "XL" and len(parts) > 1:
                    # If it's an Excel range reference
                    if ":" in parts[1] and "!" not in parts[1].split(":", 1)[1]:
                        try:
                            # Get paragraph text and position
                            orig_text = paragraph.text
                            start_pos = match.start()
                            end_pos = match.end()
                            
                            # Parse the keyword to potentially create a table
                            parser.set_word_document(doc)
                            replacement = parser.parse(keyword)
                            
                            # If a table was inserted
                            if replacement == "[TABLE_INSERTED]":
                                # Update the paragraph text to remove the keyword
                                if start_pos == 0 and end_pos == len(orig_text):
                                    # If keyword is the entire paragraph, empty it
                                    paragraph.text = ""
                                else:
                                    # Otherwise remove just the keyword
                                    paragraph.text = orig_text[:start_pos] + orig_text[end_pos:]
                            else:
                                # If table insertion failed, replace with text normally
                                paragraph.text = orig_text[:start_pos] + str(replacement) + orig_text[end_pos:]
                        except Exception as e:
                            # Handle errors
                            replacement = f"[Error: {str(e)}]"
                            paragraph.text = paragraph.text[:match.start()] + replacement + paragraph.text[match.end():]
                        
                        # Update progress
                        processed_count += 1
                        progress_bar.progress(processed_count / total_keywords)
                        progress_text.text(f"Processing keywords: {processed_count}/{total_keywords}")
                        continue
                
                # Process non-range keywords normally
                if content.split(":", 1)[0].strip().upper() == "INPUT" and keyword in input_values:
                    replacement = input_values[keyword]
                else:
                    parser.set_word_document(None)  # Don't use direct table insertion
                    replacement = parser.parse(keyword)
                
                # Replace the keyword in the paragraph text
                paragraph.text = paragraph.text[:match.start()] + str(replacement) + paragraph.text[match.end():]
                
                # Update progress
                processed_count += 1
                progress_bar.progress(processed_count / total_keywords)
                progress_text.text(f"Processing keywords: {processed_count}/{total_keywords}")
        else:
            # This is a regular paragraph, process normally
            matches = list(re.finditer(pattern, paragraph.text))
            for match in reversed(matches):  # Process in reverse to avoid index issues
                keyword = match.group(0)  # Full keyword with {{}}
                content = match.group(1)  # Content inside {{}}
                
                # If it's an INPUT keyword, use the value from our form
                if content.split(":", 1)[0].strip().upper() == "INPUT" and keyword in input_values:
                    replacement = input_values[keyword]
                else:
                    # Otherwise parse the keyword normally
                    parser.set_word_document(None)  # Don't use direct table insertion
                    replacement = parser.parse(keyword)
                
                # Replace the keyword in the paragraph text
                paragraph.text = paragraph.text[:match.start()] + str(replacement) + paragraph.text[match.end():]
                
                # Update progress
                processed_count += 1
                progress_bar.progress(processed_count / total_keywords)
                progress_text.text(f"Processing keywords: {processed_count}/{total_keywords}")
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Find all keywords in the paragraph
                    matches = list(re.finditer(pattern, paragraph.text))
                    
                    # Process each keyword
                    for match in reversed(matches):  # Process in reverse to avoid index issues
                        keyword = match.group(0)  # Full keyword with {{}}
                        content = match.group(1)  # Content inside {{}}
                        
                        # If it's an INPUT keyword, use the value from our form
                        if content.split(":", 1)[0].strip().upper() == "INPUT" and keyword in input_values:
                            replacement = input_values[keyword]
                        else:
                            # For tables, don't use direct table insertion since it would be nested
                            parser.set_word_document(None)
                            replacement = parser.parse(keyword)
                        
                        # Replace the keyword in the paragraph text
                        paragraph.text = paragraph.text[:match.start()] + str(replacement) + paragraph.text[match.end():]
                        
                        # Update progress
                        processed_count += 1
                        progress_bar.progress(processed_count / total_keywords)
                        progress_text.text(f"Processing keywords: {processed_count}/{total_keywords}")
    
    # Ensure progress is complete
    progress_bar.progress(1.0)
    progress_text.text(f"Processed {processed_count} keywords.")
    
    # Close Excel manager
    excel_mgr.close()
    
    return doc, processed_count

def main():
    st.title("Document Keyword Parser")
    st.write("Upload a Word document and an Excel spreadsheet to replace keywords in the Word document.")
    
    # File upload section
    col1, col2 = st.columns(2)
    
    with col1:
        doc_file = st.file_uploader("Upload Word Document (.docx)", type=["docx"])
    
    with col2:
        excel_file = st.file_uploader("Upload Excel Spreadsheet (.xlsx)", type=["xlsx"])
    
    # Show keyword help
    with st.expander("Keyword Reference Guide"):
        parser = keywordParser()
        st.markdown(parser.get_keyword_help())
    
    # Process the documents when both are uploaded
    if doc_file and excel_file:
        st.subheader("Processing Document")
        
        # Save uploaded files to temporary location
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_doc:
            tmp_doc.write(doc_file.getvalue())
            doc_path = tmp_doc.name
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_excel:
            tmp_excel.write(excel_file.getvalue())
            excel_path = tmp_excel.name
        
        try:
            # Process the document
            processed_doc, count = process_word_doc(doc_path, excel_path)
            
            if count > 0:
                # Save the processed document
                tmp_folder = "tmp"
                if not os.path.exists(tmp_folder):
                    os.makedirs(tmp_folder)
                output_path = os.path.join(tmp_folder, "processed_document.docx")
                processed_doc.save(output_path)
                
                # Provide download link
                with open(output_path, "rb") as file:
                    st.download_button(
                        label="Download Processed Document",
                        data=file,
                        file_name="processed_document.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                st.success(f"Successfully processed {count} keywords!")
            else:
                st.info("No keywords were processed. The document remains unchanged.")
                
        except Exception as e:
            st.error(f"An error occurred during processing: {str(e)}")
        
        finally:
            # Clean up temporary files
            os.unlink(doc_path)
            os.unlink(excel_path)
    
    # Additional information
    st.markdown("---")
    st.write("This app processes keywords in Word documents and replaces them with values from Excel.")

if __name__ == "__main__":
    main()
    