# main.py
import streamlit as st
import os
import re
import docx
import tempfile
import time
from excel_manager import excelManager # Assuming excel_manager.py is in the same directory
from keyword_parser import keywordParser # Assuming keyword_parser.py is in the same directory
from collections import Counter

def preprocess_word_doc(doc_path):
    """
    Analyze a Word document to determine what keywords it contains, using '!' separator.

    Args:
        doc_path: Path to the Word document

    Returns:
        Dictionary with keyword counts and whether Excel file is needed
    """
    doc = docx.Document(doc_path)
    pattern = r'{{(.*?)}}'

    keywords = {
        "excel": {"CELL": [], "LAST": [], "RANGE": [], "COLUMN": [], "OTHER": []},
        "input": {"text": [], "area": [], "date": [], "select": [], "check": []},
        "template": [],
        "json": [],
        "other": []
    }
    needs_excel = False
    total_keywords = 0

    def categorize_keyword(content):
        nonlocal needs_excel
        parts = content.split("!", 1) # Use '!' separator
        keyword_type = parts[0].strip().upper()

        if not keyword_type: return # Ignore empty keywords {{}}

        if keyword_type == "XL":
            needs_excel = True
            if len(parts) > 1:
                sub_parts = parts[1].split("!", 1)
                xl_subtype = sub_parts[0].strip().upper()
                if xl_subtype in keywords["excel"]:
                    keywords["excel"][xl_subtype].append(content)
                else:
                     # If subtype unknown, check if it looks like an old format/named range
                     if ':' not in parts[1] and '!' not in parts[1]: # Likely named range or old cell ref
                          keywords["excel"]["RANGE"].append(content) # Assume RANGE for named range
                     else:
                          keywords["excel"]["OTHER"].append(content) # Potentially old or invalid format
            else:
                 keywords["excel"]["OTHER"].append(content) # Invalid XL format {{XL}}

        elif keyword_type == "INPUT":
            if len(parts) > 1:
                input_parts = parts[1].split("!")
                input_type = input_parts[0].lower() if input_parts else "text"
                if input_type in keywords["input"]:
                    keywords["input"][input_type].append(content)
                else:
                    keywords["input"]["text"].append(content)
            else:
                 keywords["input"]["text"].append(content) # {{INPUT}} defaults to text

        elif keyword_type == "TEMPLATE":
            keywords["template"].append(content)
        elif keyword_type == "JSON":
            keywords["json"].append(content)
        else:
             # If not a recognized type, check if it might be an Excel named range
             if '!' not in content and ':' not in content:
                  needs_excel = True
                  keywords["excel"]["RANGE"].append(content) # Treat as potential named range
             else:
                  keywords["other"].append(content)

    # Scan paragraphs
    for paragraph in doc.paragraphs:
        matches = list(re.finditer(pattern, paragraph.text))
        total_keywords += len(matches)
        for match in matches:
            categorize_keyword(match.group(1))

    # Scan tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    matches = list(re.finditer(pattern, paragraph.text))
                    total_keywords += len(matches)
                    for match in matches:
                        categorize_keyword(match.group(1))

    summary = {
        "total_keywords": total_keywords,
        "excel_counts": {k: len(v) for k, v in keywords["excel"].items()},
        "input_counts": {k: len(v) for k, v in keywords["input"].items()},
        "template_count": len(keywords["template"]),
        "json_count": len(keywords["json"]),
        "other_count": len(keywords["other"]),
        "needs_excel": needs_excel,
        "keywords": keywords
    }
    return summary


def process_word_doc(doc_path, excel_path=None, parser=None):
    """
    Process a Word document, replacing keywords with values using the provided parser.

    Args:
        doc_path: Path to the Word document
        excel_path: Path to the Excel spreadsheet (optional - manager passed via parser)
        parser: An initialized keywordParser instance

    Returns:
        Processed document object and a count of replaced keywords
    """
    if not parser:
        raise ValueError("KeywordParser instance is required.")

    doc = docx.Document(doc_path)
    parser.set_word_document(doc) # Ensure parser has the correct document object

    pattern = r'{{(.*?)}}'
    total_keywords_initial = 0

    # Count initial keywords
    elements_to_scan = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                elements_to_scan.extend(cell.paragraphs)

    for paragraph in elements_to_scan:
        total_keywords_initial += len(re.findall(pattern, paragraph.text))

    if total_keywords_initial == 0:
        st.warning("No keywords found in the document.")
        return doc, 0

    progress_bar = st.progress(0)
    progress_text = st.empty()
    progress_text.text("Processing keywords...")

    # Process paragraph by paragraph, letting the parser handle replacements
    processed_keywords_count = 0
    elements_processed = 0
    total_elements = len(elements_to_scan)

    for paragraph in elements_to_scan:
        original_text = paragraph.text
        keywords_in_para = len(re.findall(pattern, original_text))

        if keywords_in_para > 0:
            try:
                # parser.parse will handle replacements, including potential table insertion
                processed_text = parser.parse(original_text)

                # Handle table insertion placeholder
                if "[TABLE_INSERTED]" in processed_text:
                    # Check if the keyword was the only content (strip spaces for check)
                    is_only_keyword = False
                    matches = list(re.finditer(pattern, original_text))
                    if len(matches) == 1 and matches[0].group(0).strip() == original_text.strip():
                         is_only_keyword = True

                    if is_only_keyword:
                         paragraph.text = "" # Clear paragraph if only table keyword was present
                    else:
                         # Remove placeholder but keep other text
                         paragraph.text = processed_text.replace("[TABLE_INSERTED]", "").strip()
                elif processed_text != original_text:
                    paragraph.text = processed_text

                # Estimate progress - count keywords *remaining* after parse
                keywords_remaining = len(re.findall(pattern, paragraph.text))
                processed_in_step = keywords_in_para - keywords_remaining
                processed_keywords_count += processed_in_step

            except Exception as e:
                st.error(f"Error processing content '{original_text[:50]}...': {str(e)}")
                # Keep original text on error

        elements_processed += 1
        progress = elements_processed / total_elements if total_elements > 0 else 0
        progress_bar.progress(progress)
        # Update text based on estimated keywords processed vs initial total
        progress_text.text(f"Processing: {processed_keywords_count}/{total_keywords_initial} keywords estimated...")


    progress_bar.progress(1.0)
    progress_text.text(f"Processing finished. Approximately {processed_keywords_count} keywords processed.")

    # Clean up doc reference in parser? Maybe not necessary if parser is short-lived
    # parser.set_word_document(None)

    return doc, processed_keywords_count


def display_keyword_summary(summary):
    """Display analysis summary with updated Excel categories."""
    st.write(f"Total keywords found: **{summary['total_keywords']}**")
    with st.expander("Document Analysis Summary"):
        col1, col2, col3, col4 = st.columns(4)

        with col1:
            st.markdown("**Excel Keywords (`XL!`)**")
            total_excel = sum(summary["excel_counts"].values())
            st.write(f"Total: {total_excel}")
            if summary["needs_excel"]:
                st.write("*Excel file required*")
            for subtype, count in summary["excel_counts"].items():
                 if count > 0: st.write(f"- {subtype}: {count}")

        with col2:
            st.markdown("**Input Keywords (`INPUT!`)**")
            total_inputs = sum(summary["input_counts"].values())
            st.write(f"Total: {total_inputs}")
            for input_type, count in summary["input_counts"].items():
                 if count > 0: st.write(f"- {input_type}: {count}")

        with col3:
            st.markdown("**Template Keywords (`TEMPLATE!`)**")
            st.write(f"Total: {summary['template_count']}")
            st.markdown("**JSON Keywords (`JSON!`)**")
            st.write(f"Total: {summary['json_count']}")

        with col4:
             st.markdown("**Other/Invalid**")
             st.write(f"Total: {summary['other_count']}")
             if summary['other_count'] > 0 and 'keywords' in summary and summary['keywords']['other']:
                  st.caption("Examples:")
                  for item in summary['keywords']['other'][:3]: # Show first few
                       st.caption(f"`{{{{{item}}}}}`")


def main():
    st.title("Document Keyword Parser")
    st.markdown("""
    Upload a Word document (`.docx`) containing keywords (e.g., `{{XL!CELL!A1}}`, `{{INPUT!text!Name}}`).
    The tool analyzes keywords, prompts for required files (like Excel), gathers user input if needed,
    processes the document, and provides a download link. Uses `!` as keyword separator.
    """)

    # Initialize parser instance for help text display
    if 'keyword_parser_instance_for_help' not in st.session_state:
         st.session_state.keyword_parser_instance_for_help = keywordParser()
    with st.expander("Keyword Reference Guide"):
        st.markdown(st.session_state.keyword_parser_instance_for_help.get_keyword_help())

    # --- State Management ---
    # Initialize state variables if they don't exist
    default_state = {
        'doc_uploaded': False, 'doc_path': None, 'analysis_summary': None,
        'excel_uploaded': False, 'excel_path': None, 'excel_manager_instance': None,
        'keyword_parser_instance': None, 'form_submitted_main': False, 'input_values_main': {},
        'processing_started': False, 'processed_doc_path': None, 'processed_count': 0
    }
    for key, value in default_state.items():
        if key not in st.session_state:
            st.session_state[key] = value

    # --- Reset Button ---
    if st.sidebar.button("Reset Application State"):
        # Clean up temp files
        if st.session_state.doc_path and os.path.exists(st.session_state.doc_path): os.unlink(st.session_state.doc_path)
        if st.session_state.excel_path and os.path.exists(st.session_state.excel_path): os.unlink(st.session_state.excel_path)
        if st.session_state.processed_doc_path and os.path.exists(st.session_state.processed_doc_path): os.unlink(st.session_state.processed_doc_path)
        # Close Excel Manager if open
        if st.session_state.excel_manager_instance: st.session_state.excel_manager_instance.close()
        # Reset state variables
        for key in default_state:
            st.session_state[key] = default_state[key]
        st.rerun()

    # --- Step 1: Upload Word Document ---
    st.header("Step 1: Upload Document")
    doc_file = st.file_uploader("Upload Word Document (.docx)", type=["docx"], key="main_doc_uploader")

    if doc_file and not st.session_state.doc_uploaded:
         # Reset relevant states for new upload
        st.session_state.update({k: v for k, v in default_state.items() if k not in ['keyword_parser_instance_for_help']}) # Keep help parser
        # Save new doc
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_doc:
            tmp_doc.write(doc_file.getvalue())
            st.session_state.doc_path = tmp_doc.name
        st.session_state.doc_uploaded = True
        st.rerun()

    # --- Step 2: Analysis & Conditional Excel Upload ---
    if st.session_state.doc_uploaded:
        st.header("Step 2: Analysis & File Uploads")
        if not st.session_state.analysis_summary:
            st.info("Analyzing document...")
            try:
                summary = preprocess_word_doc(st.session_state.doc_path)
                st.session_state.analysis_summary = summary
                st.rerun()
            except Exception as e:
                st.error(f"Analysis failed: {e}")
                st.session_state.doc_uploaded = False # Allow re-upload

        if st.session_state.analysis_summary:
            display_keyword_summary(st.session_state.analysis_summary)
            needs_excel = st.session_state.analysis_summary["needs_excel"]

            excel_file = None
            if needs_excel:
                excel_file = st.file_uploader("Upload Required Excel Spreadsheet (.xlsx)", type=["xlsx"], key="main_excel_uploader")
                if excel_file and not st.session_state.excel_uploaded:
                    # Save new excel file
                    if st.session_state.excel_path and os.path.exists(st.session_state.excel_path): os.unlink(st.session_state.excel_path) # Clean old temp excel
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_excel:
                        tmp_excel.write(excel_file.getvalue())
                        st.session_state.excel_path = tmp_excel.name
                    st.session_state.excel_uploaded = True
                    # Reset excel manager instance as file changed
                    if st.session_state.excel_manager_instance: st.session_state.excel_manager_instance.close()
                    st.session_state.excel_manager_instance = None
                    st.rerun()
                elif not excel_file and not st.session_state.excel_uploaded:
                     st.warning("Excel file is required based on document analysis.")


            # --- Initialize Managers (only once per valid file state) ---
            if needs_excel and st.session_state.excel_path and not st.session_state.excel_manager_instance:
                 try:
                      st.session_state.excel_manager_instance = excelManager(st.session_state.excel_path)
                 except Exception as e:
                      st.error(f"Failed to load Excel file: {e}")
                      st.session_state.excel_uploaded = False # Reset upload status

            # Always ensure parser instance exists, update if excel manager changes
            current_excel_manager = st.session_state.excel_manager_instance if needs_excel else None
            if not st.session_state.keyword_parser_instance or st.session_state.keyword_parser_instance.excel_manager != current_excel_manager:
                 st.session_state.keyword_parser_instance = keywordParser(current_excel_manager)


            # --- Step 3: User Input Form (if needed) ---
            parser = st.session_state.keyword_parser_instance
            # Check for input keywords using the analysis summary
            has_inputs = sum(st.session_state.analysis_summary['input_counts'].values()) > 0

            if has_inputs:
                 st.header("Step 3: Provide User Inputs")
                 if not st.session_state.form_submitted_main:
                      with st.form(key="main_input_form"):
                           temp_inputs = {}
                           # Use analysis summary to find all input keywords (more reliable than parser state)
                           all_input_keywords = [item for sublist in st.session_state.analysis_summary['keywords']['input'].values() for item in sublist]
                           unique_input_contents = sorted(list(set(all_input_keywords))) # Get unique input definitions

                           for content in unique_input_contents:
                                # Create field using parser's helper function (doesn't store value in parser here)
                                # Key must be unique and stable for Streamlit form state
                                field_key = f"form_{content}"
                                temp_inputs[content] = parser._create_input_field(content) # Use unique key inside

                           submitted = st.form_submit_button("Submit Inputs")
                           if submitted:
                                # Store the submitted values based on their content definition
                                st.session_state.input_values_main = {content: st.session_state.get(f"form_{content}", "") for content in unique_input_contents}
                                st.session_state.form_submitted_main = True
                                # Update the parser's internal values for the processing step
                                parser.input_values = {}
                                # Map stored values back to potential keyword occurrences (this is tricky, assumes simple replacement)
                                # A better approach might be needed if the same INPUT definition appears multiple times
                                for content, value in st.session_state.input_values_main.items():
                                     parser.input_values[f"{{{{{content}}}}}"] = value

                                st.success("Inputs submitted.")
                                st.rerun() # Rerun to proceed to processing step
                 else:
                      st.success("Inputs already submitted.")
                      # Display submitted values read-only? Optional.


            # --- Step 4: Process ---
            st.header("Step 4: Process Document")
            # Determine if ready to process
            ready_to_process = st.session_state.doc_uploaded and \
                               (not needs_excel or st.session_state.excel_uploaded) and \
                               (not has_inputs or st.session_state.form_submitted_main)

            process_button_disabled = not ready_to_process or st.session_state.processing_started

            if st.button("Process Document", disabled=process_button_disabled, key="main_process_btn"):
                st.session_state.processing_started = True
                st.session_state.processed_doc_path = None # Clear previous

                st.info("Processing document... This may take a moment.")
                try:
                     # Ensure parser has the submitted inputs
                     parser = st.session_state.keyword_parser_instance
                     if has_inputs and st.session_state.form_submitted_main:
                         # Re-apply submitted values to parser instance before processing
                          parser.input_values = {}
                          for content, value in st.session_state.input_values_main.items():
                               parser.input_values[f"{{{{{content}}}}}"] = value

                     # Process
                     processed_doc, count = process_word_doc(
                          st.session_state.doc_path,
                          st.session_state.excel_path,
                          parser=parser
                     )

                     if processed_doc:
                          tmp_folder = "tmp_processed_main"
                          if not os.path.exists(tmp_folder): os.makedirs(tmp_folder)
                          # Use original filename for output
                          base_name = os.path.basename(st.session_state.doc_path)
                          output_filename = f"processed_{base_name}" if not base_name.startswith("tmp") else "processed_document.docx"
                          output_path = os.path.join(tmp_folder, output_filename)

                          processed_doc.save(output_path)
                          st.session_state.processed_doc_path = output_path
                          st.session_state.processed_count = count
                          st.success(f"Processing Complete! Approximately {count} keywords processed.")
                          # Rerun needed to show download button correctly after processing finishes
                          st.rerun()
                     else:
                           st.warning("Processing did not return a document.")


                except Exception as e:
                     st.error(f"Error during processing: {e}")
                finally:
                      # Close excel manager instance if it exists
                     if st.session_state.excel_manager_instance:
                          st.session_state.excel_manager_instance.close()
                          st.session_state.excel_manager_instance = None
                     st.session_state.processing_started = False # Reset processing flag


            # --- Step 5: Download ---
            if st.session_state.processed_doc_path:
                 st.header("Step 5: Download Result")
                 st.success(f"Document processed. Approximately {st.session_state.processed_count} keywords replaced.")
                 try:
                     with open(st.session_state.processed_doc_path, "rb") as fp:
                          st.download_button(
                               label="Download Processed Document",
                               data=fp,
                               file_name=os.path.basename(st.session_state.processed_doc_path),
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                          )
                 except FileNotFoundError:
                      st.error("Processed file not found. Please try processing again.")
                      st.session_state.processed_doc_path = None # Reset path


if __name__ == "__main__":
    main()