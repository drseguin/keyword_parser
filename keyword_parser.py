# keyword_parser.py
import re
import json
import os
import streamlit as st
from datetime import date, datetime
from excel_manager import excelManager
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

class keywordParser:
    """
    A parser class that processes various keywords and extracts data from Excel,
    handles user input, and processes templates and JSON data using '!' as a separator.
    """

    def __init__(self, excel_manager=None):
        """
        Initialize the keyword parser.

        Args:
            excel_manager: An instance of excelManager to use for Excel operations.
                           If None, a new instance will be created when needed.
        """
        self.excel_manager = excel_manager
        self.pattern = r'{{(.*?)}}'
        self.has_input_fields = False
        self.form_submitted = False
        self.word_document = None
        self.input_values = {}  # Store input values

    def set_word_document(self, doc):
        """Set the word document for direct table insertion."""
        self.word_document = doc

    def parse(self, input_string):
        """
        Parse input string and process any keywords found.

        Args:
            input_string: The string containing keywords to parse.

        Returns:
            Processed string with keywords replaced with their values.
        """
        if not input_string:
            return input_string

        # Find all keywords in the input string
        matches = list(re.finditer(self.pattern, input_string))

        # First handle all INPUT keywords
        input_keywords = []
        for match in matches:
            content = match.group(1)  # Content inside {{}}
            keyword = match.group(0)  # The full {{keyword}}
            parts = content.split("!", 1) # Use '!' as separator
            keyword_type = parts[0].strip().upper()

            if keyword_type == "INPUT":
                input_keywords.append((keyword, content))

        # If we have input fields, process them first using a form
        if input_keywords and not self.form_submitted:
            with st.form(key=f"input_form_{id(input_string)}"):
                st.subheader("Please provide input values:")

                # Create input fields and store their values
                temp_input_values = {}
                for keyword, content in input_keywords:
                    value = self._create_input_field(content)
                    temp_input_values[keyword] = value

                # Add submit button
                submit = st.form_submit_button("Submit")
                if submit:
                    self.input_values.update(temp_input_values) # Store values upon submission
                    self.form_submitted = True
                    st.rerun() # Rerun to process the rest of the keywords
                else:
                    # If not submitted, show message and don't process further yet
                    st.stop()


        # After processing inputs or if no inputs, process all keywords
        result = input_string
        for match in matches:
            keyword = match.group(0)  # Full keyword with {{}}
            content = match.group(1)  # Content inside {{}}

            # If this is an INPUT keyword we've already processed
            if keyword in self.input_values:
                replacement = self.input_values[keyword]
            else:
                replacement = self._process_keyword(content)

            # Replace the keyword with its value
            # Ensure replacement is string, handle potential None values
            result = result.replace(keyword, str(replacement) if replacement is not None else "", 1)


        return result

    def _create_input_field(self, content):
        """
        Create an appropriate input field based on the INPUT keyword using '!' separator.

        Args:
            content: The content inside the {{ }} brackets.

        Returns:
            The value from the input field.
        """
        if not content:
            return "[Invalid input reference]"

        # Split the content into tokens using '!'
        tokens = content.split("!")
        if len(tokens) < 2:
            return "[Invalid INPUT format]"

        # Get the keyword type (INPUT) and input type (text, area, date, select, check)
        keyword_type = tokens[0].strip().upper()
        input_type = tokens[1].strip().lower() if len(tokens) > 1 else ""

        # Check for valid INPUT keyword
        if keyword_type != "INPUT":
            return "[Invalid INPUT keyword]"

        # Handle text input - {{INPUT!text!label!value}}
        if input_type == "text":
            label = tokens[2] if len(tokens) > 2 else ""
            default_value = tokens[3] if len(tokens) > 3 else ""
            return st.text_input(
                label=label,
                value=default_value,
                label_visibility="visible",
                key=f"input_field_{content}" # Unique key
            )

        # Handle text area - {{INPUT!area!label!value!height}}
        elif input_type == "area":
            label = tokens[2] if len(tokens) > 2 else ""
            default_value = tokens[3] if len(tokens) > 3 else ""
            height_px = tokens[4] if len(tokens) > 4 else None

            # Convert height to integer if provided
            height = None
            if height_px:
                try:
                    height = int(height_px)
                except ValueError:
                    # If height is not a valid integer, ignore it
                    pass

            # Set height if provided, otherwise use default
            if height:
                return st.text_area(
                    label=label,
                    value=default_value,
                    height=height,
                    label_visibility="visible",
                     key=f"input_field_{content}" # Unique key
                )
            else:
                return st.text_area(
                    label=label,
                    value=default_value,
                    label_visibility="visible",
                     key=f"input_field_{content}" # Unique key
                )

        # Handle date input - {{INPUT!date!label!value!format}}
        elif input_type == "date":
            label = tokens[2] if len(tokens) > 2 else ""
            default_value_str = tokens[3] if len(tokens) > 3 else "today"
            date_format = tokens[4] if len(tokens) > 4 else "YYYY/MM/DD"

            # Handle "today" default value
            if default_value_str.lower() == "today":
                default_date = date.today()
            else:
                try:
                    # Try to parse the date based on the format
                    if date_format == "YYYY/MM/DD":
                        default_date = datetime.strptime(default_value_str, "%Y/%m/%d").date()
                    elif date_format == "DD/MM/YYYY":
                        default_date = datetime.strptime(default_value_str, "%d/%m/%Y").date()
                    elif date_format == "MM/DD/YYYY":
                        default_date = datetime.strptime(default_value_str, "%m/%d/%Y").date()
                    else:
                        # Default to ISO format if format is not recognized
                        default_date = datetime.strptime(default_value_str, "%Y-%m-%d").date()
                except ValueError:
                    default_date = date.today()

            date_value = st.date_input(
                label=label,
                value=default_date,
                label_visibility="visible",
                 key=f"input_field_{content}" # Unique key
            )

            # Return the date in the requested format
            if date_format == "YYYY/MM/DD":
                return date_value.strftime("%Y/%m/%d")
            elif date_format == "DD/MM/YYYY":
                return date_value.strftime("%d/%m/%Y")
            elif date_format == "MM/DD/YYYY":
                return date_value.strftime("%m/%d/%Y")
            else:
                return date_value.strftime("%Y/%m/%d")  # Default format

        # Handle select box - {{INPUT!select!label!options}}
        elif input_type == "select":
            label = tokens[2] if len(tokens) > 2 else ""
            options_str = tokens[3] if len(tokens) > 3 else ""

            # Parse options (comma-separated)
            options = [opt.strip() for opt in options_str.split(",")] if options_str else []

            if not options:
                return "[No options provided]"

            return st.selectbox(
                label=label,
                options=options,
                label_visibility="visible",
                 key=f"input_field_{content}" # Unique key
            )

        # Handle checkbox - {{INPUT!check!label!value}}
        elif input_type == "check":
            label = tokens[2] if len(tokens) > 2 else ""
            default_value_str = tokens[3].lower() if len(tokens) > 3 else "false"

            # Convert string value to boolean
            default_value = default_value_str == "true"

            return st.checkbox(
                label=label,
                value=default_value,
                label_visibility="visible",
                 key=f"input_field_{content}" # Unique key
            )

        # Default for unrecognized input types
        else:
            return f"[Unsupported input type: {input_type}]"

    def _process_keyword(self, content):
        """
        Process a single keyword content and return the corresponding value using '!' separator.

        Args:
            content: The content inside the {{ }} brackets.

        Returns:
            The processed value of the keyword.
        """
        parts = content.split("!", 1) # Use '!' as separator
        keyword_type = parts[0].strip().upper()

        # Process Excel data keywords
        if keyword_type == "XL":
            return self._process_excel_keyword(parts[1] if len(parts) > 1 else "")

        # Process user input keywords - these should already be handled in parse()
        elif keyword_type == "INPUT":
            # Fallback if not handled by the form (e.g., in tester_app without form)
             params = parts[1] if len(parts) > 1 else ""
             return self._process_input_keyword(params)


        # Process template keywords
        elif keyword_type == "TEMPLATE":
            return self._process_template_keyword(parts[1] if len(parts) > 1 else "")

        # Process JSON keywords
        elif keyword_type == "JSON":
            return self._process_json_keyword(parts[1] if len(parts) > 1 else "")

        # Unknown keyword type
        else:
            # If no known keyword type, treat it as a potential named range for XL
             return self._process_excel_keyword(f"RANGE!{content}")
            # return f"[Unknown keyword type: {keyword_type}]"

    def _process_input_keyword(self, params):
        """Process INPUT keywords directly if needed (fallback). Uses '!' separator."""
        input_parts = params.split("!") # Use '!' separator
        input_type = input_parts[0].lower() if input_parts else ""

        if input_type == "text" or input_type == "area":
            label = input_parts[1] if len(input_parts) > 1 else ""
            default_value = input_parts[2] if len(input_parts) > 2 else ""
            return default_value

        elif input_type == "date":
            # Use already imported datetime modules correctly
            today = date.today()
            return today.strftime("%Y/%m/%d")

        elif input_type == "select":
            options_str = input_parts[2] if len(input_parts) > 2 else ""
            options = [opt.strip() for opt in options_str.split(",")] if options_str else []
            return options[0] if options else ""

        elif input_type == "check":
            default_value_str = input_parts[2].lower() if len(input_parts) > 2 else "false"
            return default_value_str == "true"

        else:
            return params if params else "[Input value]"

    def _process_excel_keyword(self, content):
        """Process Excel-related keywords with new structure and '!' separator."""
        if not content:
            return "[Invalid Excel reference]"

        if not self.excel_manager:
            return "[Excel manager not initialized]"

        parts = content.split("!") # Use '!' as separator
        if len(parts) < 2:
             # Attempt to handle old format or named range as RANGE
            if ':' in content: # Could be old range XL:Sheet!A1:B2 or XL:A1:B2
                 if '!' in content.split(':')[0]: # Old range with sheet XL:Sheet!A1:B2
                     sheet_ref, cell_range = content.split('!', 1)
                     return self._call_excel_method("RANGE", f"{sheet_ref}!{cell_range}")
                 else: # Old range without sheet XL:A1:B2
                     return self._call_excel_method("RANGE", content)
            elif content.startswith(':'): # Old LAST format XL::A1 or XL::Sheet!A1
                 return self._call_excel_method("LAST", content[1:]) # Remove leading ':'
            else: # Assume it's a named range or old cell format XL:A1 or XL:Sheet!A1
                if '!' in content: # Old cell with sheet XL:Sheet!A1
                     return self._call_excel_method("CELL", content)
                else: # Old cell without sheet XL:A1 or a named range
                    # Try as cell first, if error, treat as named range
                    try:
                       return self._call_excel_method("CELL", content)
                    except ValueError:
                       return self._call_excel_method("RANGE", content) # Treat as named range

            # return f"[Invalid XL format: {content}]"


        xl_type = parts[0].strip().upper()
        xl_params = "!".join(parts[1:]) # Rejoin remaining parts

        return self._call_excel_method(xl_type, xl_params)


    def _call_excel_method(self, xl_type, xl_params):
        """Helper function to call the appropriate excelManager method."""
        available_sheets = self.excel_manager.get_sheet_names()
        sheet_name_map = {sheet.lower(): sheet for sheet in available_sheets}

        try:
            # {{XL!CELL!A1}} or {{XL!CELL!Sheet2!B5}}
            if xl_type == "CELL":
                sheet_name, cell_ref = self._get_sheet_and_ref(xl_params, available_sheets[0], sheet_name_map)
                return self.excel_manager.read_cell(sheet_name, cell_ref)

            # {{XL!LAST!A1}} or {{XL!LAST!Sheet2!B5}}
            # {{XL!LAST!sheet_name!A1!Title}}
            elif xl_type == "LAST":
                last_parts = xl_params.split("!")
                if len(last_parts) >= 3: # Title format: {{XL!LAST!sheet_name!A1!Title}}
                    sheet_name_ref = last_parts[0]
                    cell_ref = last_parts[1]
                    title = last_parts[2]
                    actual_sheet_name = sheet_name_map.get(sheet_name_ref.lower(), sheet_name_ref) # Allow direct sheet name or lookup
                    if actual_sheet_name not in available_sheets: return f"[Sheet not found: {actual_sheet_name}]"
                    return self.excel_manager.read_title_total(actual_sheet_name, cell_ref, title)
                else: # Basic LAST format: {{XL!LAST!A1}} or {{XL!LAST!Sheet2!B5}}
                    sheet_name, cell_ref = self._get_sheet_and_ref(xl_params, available_sheets[0], sheet_name_map)
                    return self.excel_manager.read_total(sheet_name, cell_ref)

            # {{XL!RANGE!Sales!C3:C7}} or {{XL!RANGE!named_range}}
            elif xl_type == "RANGE":
                sheet_name, range_ref = self._get_sheet_and_ref(xl_params, available_sheets[0], sheet_name_map)
                # If range_ref doesn't contain ':' it's likely a named range or invalid
                if ':' not in range_ref:
                    # Attempt to read as named range (assuming excelManager handles it)
                    # Or handle named range lookup if excelManager doesn't
                     # For now, assume read_range might handle named ranges or error out.
                    pass # Continue to read_range

                data = self.excel_manager.read_range(sheet_name, range_ref)
                if self.word_document and data:
                    return self._create_word_table(data)
                else:
                    return self._format_table(data)

            # {{XL!COLUMN!Sales!A1,C1,E1}} or {{XL!COLUMN!Sales!Revenue,Expense,Profit!12}}
            elif xl_type == "COLUMN":
                col_parts = xl_params.split("!")
                if len(col_parts) < 2: return "[Invalid COLUMN format]"

                sheet_ref = col_parts[0]
                columns_input = col_parts[1].strip('"') # Cell refs or titles

                actual_sheet_name = sheet_name_map.get(sheet_ref.lower(), sheet_ref) # Allow direct sheet name or lookup
                if actual_sheet_name not in available_sheets: return f"[Sheet not found: {actual_sheet_name}]"

                start_row = None
                use_titles = False

                if len(col_parts) > 2: # Optional start row provided, implies using titles
                    try:
                        start_row = int(col_parts[2])
                        use_titles = True
                    except ValueError:
                        return "[Invalid start row for COLUMN]"
                else:
                    # Determine if using titles based on input format (heuristic: check for letters vs numbers)
                    # A more robust check might be needed, e.g., trying to parse as cell ref
                    if not any(char.isdigit() for char in columns_input.replace(',', '')):
                         use_titles = True
                         start_row = 1 # Default start row for titles if not specified
                    # else: use_titles = False (default)


                data = self.excel_manager.read_columns(actual_sheet_name, columns_input, use_titles=use_titles, start_row=start_row)

                if self.word_document and data:
                    return self._create_word_table(data)
                else:
                    return self._format_table(data)


            else:
                return f"[Unknown XL type: {xl_type}]"

        except Exception as e:
            self.excel_manager.logger.error(f"Error processing XL keyword '{content}': {str(e)}", exc_info=True)
            return f"[Error processing XL: {str(e)}]"

    def _get_sheet_and_ref(self, params, default_sheet, sheet_map):
        """Helper to extract sheet name and cell/range reference."""
        parts = params.split("!")
        if len(parts) > 1 and parts[0].strip("'").lower() in sheet_map:
            # Sheet name is explicitly provided
            sheet_name = sheet_map[parts[0].strip("'").lower()]
            reference = "!".join(parts[1:]) # Rejoin if ref itself contains '!'
        else:
            # Use default sheet
            sheet_name = default_sheet
            reference = params
        return sheet_name, reference


    def _format_table(self, data):
        """
        Format the data as a formatted table for Word or text.
        """
        if self.word_document:
            return self._create_word_table(data)

        if not data or not isinstance(data, list) or not all(isinstance(row, list) for row in data):
             return str(data) # Return raw data if not a list of lists

        # Calculate column widths
        col_widths = [0] * (max(len(row) for row in data) if data else 0)
        for row in data:
            for i, cell in enumerate(row):
                cell_str = str(cell) if cell is not None else ""
                if i < len(col_widths):
                     col_widths[i] = max(col_widths[i], len(cell_str))


        # Create the table as a string
        result = []
        for row_index, row in enumerate(data):
            row_str = []
            for i, cell in enumerate(row):
                 cell_str = str(cell) if cell is not None else ""
                 # Basic alignment (numbers right, text left) - simplistic check
                 try:
                      # Attempt to convert to float, fails for non-numeric strings
                      float(cell_str.replace(',', '').replace('$', ''))
                      formatted = cell_str.rjust(col_widths[i])
                 except (ValueError, TypeError):
                      formatted = cell_str.ljust(col_widths[i])

                 if i < len(col_widths):
                     row_str.append(formatted)


            result.append(" | ".join(row_str))

            # Add a separator after the header row (if more than one row exists)
            if row_index == 0 and len(data) > 1:
                separator = ["-" * width for width in col_widths]
                result.append("-+-".join(separator))

        return "\n".join(result)

    def _create_word_table(self, data):
        """
        Create a visually appealing table directly in the Word document with proper styling.
        """
        if not data or not isinstance(data, list) or not all(isinstance(row, list) for row in data):
            return str(data) # Return raw data representation if not table format

        num_rows = len(data)
        num_cols = max(len(row) for row in data) if num_rows > 0 else 0
        if num_cols == 0: return "[Empty Table Data]"


        # Create the table with the 'Table Grid' style for consistent borders
        table = self.word_document.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'

        # Set overall table properties for better appearance
        try:
            # Set the table to auto-fit contents
            table.autofit = True

            # Fill the table with data and apply formatting
            for i, row in enumerate(data):
                for j in range(num_cols): # Ensure all cells in the row are processed
                    cell_value = row[j] if j < len(row) else None # Handle ragged rows

                    # Format the cell value (handle None)
                    if cell_value is None:
                         cell_text = ""
                    elif isinstance(cell_value, (int, float)):
                        cell_text = f"{cell_value:,.2f}" # Format numbers nicely
                    else:
                        cell_text = str(cell_value)

                    cell = table.cell(i, j)
                    # Check if cell contains multiple paragraphs and clear extra ones
                    if len(cell.paragraphs) > 1:
                         for p in cell.paragraphs[1:]:
                              p.clear() # Remove extra default paragraphs
                    # Ensure there's at least one paragraph to write to
                    if not cell.paragraphs:
                         cell.add_paragraph()

                    run = cell.paragraphs[0].clear().add_run(cell_text) # Clear and add new run


                    # Apply consistent font size
                    run.font.size = Pt(10)

                    # Apply padding within cells (apply to paragraph format)
                    cell.paragraphs[0].paragraph_format.space_before = Pt(3)
                    cell.paragraphs[0].paragraph_format.space_after = Pt(3)

                    # Format header row (first row)
                    if i == 0:
                        run.font.bold = True
                        # Add light gray shading to header row
                        tcPr = cell._tc.get_or_add_tcPr()
                        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>')
                        tcPr.append(shading_elm)
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                         # Right-align numbers for better readability
                         # More robust check for numbers including currency
                        is_numeric = False
                        try:
                             # Try converting after removing currency symbols and commas
                             float(str(cell_value).replace('$', '').replace(',', ''))
                             is_numeric = True
                        except (ValueError, TypeError):
                             is_numeric = False

                        if is_numeric:
                             cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        else:
                             cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Apply alternating row colors (excluding header)
            for i in range(1, num_rows):
                if i % 2 != 0:  # Apply shading to odd rows (1, 3, 5...)
                    for j in range(num_cols):
                        cell = table.cell(i, j)
                        tcPr = cell._tc.get_or_add_tcPr()
                        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')  # Very light gray
                        tcPr.append(shading_elm)

        except Exception as e:
            # If enhanced formatting fails, continue with basic table
            self.excel_manager.logger.warning(f"Some table formatting could not be applied: {str(e)}")
            pass

        # Add a paragraph after the table for better spacing
        # self.word_document.add_paragraph() # Removed this as it might add too much space depending on context

        # Return a placeholder - the actual table has been added to the document
        return "[TABLE_INSERTED]"


    def _process_template_keyword(self, content):
        """Process template keywords using '!' separator."""
        if not content:
            return "[Invalid TEMPLATE reference]"

        try:
            # Split into filename and optional parameters using '!'
            parts = content.split("!")
            filename = parts[0].strip()

            # Handle library templates {{TEMPLATE!LIBRARY!template_name!version}}
            if filename.upper() == "LIBRARY":
                 if len(parts) > 1:
                     template_name = parts[1].strip() if len(parts) > 1 else ""
                     template_version = parts[2].strip() if len(parts) > 2 else "DEFAULT"
                     # Implement template library lookup here
                     return f"[Template Library: {template_name} (Version: {template_version})]"
                 return "[Invalid library template reference]"


            # Handle file-based templates
            if not os.path.exists(filename):
                return f"[Template file not found: {filename}]"

            # Read the file
            with open(filename, 'r', encoding='utf-8') as file: # Added encoding
                file_content = file.read()

            # Check for additional parameters (section, line, paragraph, vars)
            if len(parts) > 1:
                param_part = "!".join(parts[1:]) # Rejoin params in case '!' is in value

                 # Handle section/bookmark {{TEMPLATE!filename.docx!section=name}}
                if param_part.startswith("section="):
                    section_name = param_part.split("section=")[1].split(",")[0].strip()
                    # Implement section extraction logic here
                    return f"[Section {section_name} from {filename}]"

                # Handle specific line {{TEMPLATE!filename.txt!line=5}}
                elif param_part.startswith("line="):
                    try:
                        line_number = int(param_part.split("line=")[1].split(",")[0].strip())
                        lines = file_content.splitlines()
                        if 0 <= line_number - 1 < len(lines): # Adjust for 0-based index
                            return lines[line_number - 1]
                        return f"[Line {line_number} not found in {filename}]"
                    except (ValueError, IndexError):
                         return f"[Invalid line number in {param_part}]"


                # Handle specific paragraph {{TEMPLATE!filename.docx!paragraph=3}}
                elif param_part.startswith("paragraph="):
                    try:
                         para_number = int(param_part.split("paragraph=")[1].split(",")[0].strip())
                         # Simple paragraph split (might need refinement based on docx structure)
                         paragraphs = file_content.split("\n\n")
                         if 0 <= para_number - 1 < len(paragraphs): # Adjust for 0-based index
                             return paragraphs[para_number - 1]
                         return f"[Paragraph {para_number} not found in {filename}]"
                    except (ValueError, IndexError):
                         return f"[Invalid paragraph number in {param_part}]"


                # Handle variable substitution {{TEMPLATE!filename.docx!VARS(name=John,date=2025-04-01)}}
                elif param_part.startswith("VARS("):
                    try:
                         vars_text = param_part.split("VARS(")[1].split(")")[0]
                         var_pairs = vars_text.split(",")

                         # Create a dictionary of variables
                         variables = {}
                         for pair in var_pairs:
                             if "=" in pair:
                                 key, value = pair.split("=", 1)
                                 # Recursively parse value if it's a keyword
                                 variables[key.strip()] = self.parse(value.strip())


                         # Replace variables in the template
                         result = file_content
                         for key, value in variables.items():
                             result = result.replace(f"{{{key}}}", str(value)) # Ensure value is string

                         return result
                    except IndexError:
                         return f"[Invalid VARS format in {param_part}]"


            # Return the entire file content if no specific parameters
            return file_content

        except Exception as e:
            self.excel_manager.logger.error(f"Error processing TEMPLATE keyword '{content}': {str(e)}", exc_info=True)
            return f"[Error in TEMPLATE: {str(e)}]"


    def _process_json_keyword(self, content):
        """Process JSON keywords using '!' separator."""
        if not content:
            return "[Invalid JSON reference]"

        try:
             # Split into filename, path, and optional transformation using '!'
            parts = content.split("!")
            if len(parts) < 2: return "[Invalid JSON format: Filename and path required]"

            filename = parts[0].strip()
            json_path = parts[1].strip()
            transform_type = parts[2].strip().upper() if len(parts) > 2 else None


            # Check if filename is from another reference
            if filename.startswith("{{") and filename.endswith("}}"):
                # Recursively parse the reference
                filename = self.parse(filename)

            # Check if file exists
            if not os.path.exists(filename):
                return f"[JSON file not found: {filename}]"

            # Read the JSON file
            with open(filename, 'r', encoding='utf-8') as file: # Added encoding
                json_data = json.load(file)

            # Simplistic JSONPath implementation (needs a library for full support)
            if json_path.startswith("$."):
                path_parts = json_path[2:].split(".")
                current = json_data

                for part in path_parts:
                    # Handle array indexing like array[0] or [*]
                    if "[" in part and part.endswith("]"):
                        key = part.split("[")[0]
                        index_str = part.split("[")[1][:-1]

                        # Handle accessing the array itself if key is empty
                        if key:
                            if key not in current: return f"[JSON key not found: {key}]"
                            current = current[key]
                            if not isinstance(current, list): return f"[JSON path error: {key} is not an array]"

                        # Handle index or wildcard
                        if index_str == '*':
                             # This simplistic implementation doesn't fully support complex [*] behavior
                             # It might just return the list itself or error if used mid-path incorrectly
                             # A proper JSONPath library is needed for full support
                             pass # 'current' remains the list for now
                        else:
                             try:
                                 index = int(index_str)
                                 if index >= len(current): return f"[JSON index out of bounds: {index}]"
                                 current = current[index]
                             except (ValueError, IndexError, TypeError):
                                 return f"[Invalid JSON array index: {index_str}]"
                    else:
                        # Handle dynamic property names using keywords
                        if part.startswith("{{") and part.endswith("}}"):
                            part = self.parse(part) # Recursively parse key

                        if not isinstance(current, dict) or part not in current:
                             return f"[JSON key not found: {part}]"
                        current = current[part]


                 # Check for transformations if specified as the third part
                if transform_type:
                    if transform_type == "SUM" and isinstance(current, list):
                        try:
                            # Attempt to sum, converting elements to float
                             return sum(float(str(x).replace(',','').replace('$','')) for x in current if x is not None)
                        except (ValueError, TypeError):
                            return f"[Cannot SUM non-numeric values in list]"

                    elif transform_type.startswith("JOIN(") and transform_type.endswith(")"):
                        delimiter = transform_type[5:-1]
                        if isinstance(current, list):
                            return delimiter.join(str(x) for x in current if x is not None)
                        return str(current) # Join on single item returns the item as string


                    elif transform_type.startswith("BOOL(") and transform_type.endswith(")"):
                         yes_no = transform_type[5:-1].split("/")
                         yes_text = yes_no[0] if len(yes_no) > 0 else "Yes"
                         no_text = yes_no[1] if len(yes_no) > 1 else "No"

                         # Handle boolean conversion robustly
                         bool_value = False
                         if isinstance(current, bool):
                              bool_value = current
                         elif isinstance(current, str):
                              bool_value = current.lower() in ['true', 'yes', '1', 'on']
                         elif isinstance(current, (int, float)):
                              bool_value = current != 0

                         return yes_text if bool_value else no_text

                # Return the final value if no transformation or if transformation failed
                return current


            else:
                 return f"[Invalid JSONPath (must start with $.): {json_path}]"


        except json.JSONDecodeError:
            return f"[Error decoding JSON file: {filename}]"
        except Exception as e:
            self.excel_manager.logger.error(f"Error processing JSON keyword '{content}': {str(e)}", exc_info=True)
            return f"[Error in JSON: {str(e)}]"


    def reset_form_state(self):
        """Reset the form submission state and clear cached values."""
        self.form_submitted = False
        self.input_values = {}

    def clear_input_cache(self):
        """Clear the cached user inputs stored in session state (for Streamlit apps)."""
        # Find all INPUT keys in session state derived from input fields and clear them
        keys_to_clear = [key for key in st.session_state.keys() if key.startswith("input_field_INPUT!")]
        for key in keys_to_clear:
            # Reset based on type - needs knowledge of type or safe reset
             if isinstance(st.session_state[key], bool):
                 st.session_state[key] = False # Default for checkbox
             elif isinstance(st.session_state[key], (int, float)):
                  st.session_state[key] = 0 # Default for number (if ever used)
             else:
                 st.session_state[key] = "" # Default for text/area/select/date


        # Also clear the parser's internal state
        self.reset_form_state()



    def get_keyword_help(self):
        """
        Get help text explaining how to use keywords with '!' separator.

        Returns:
            A string with help information about available keywords.
        """
        help_text = """
        ## Keyword System Help (using '!' separator)

        ### Excel Data Keywords (`{{XL!...}}`)
        Keywords to fetch data from an Excel file.

        * **`{{XL!CELL!cell_ref}}`**: Get value from a single cell.
            * Example: `{{XL!CELL!A1}}`
            * Example with sheet: `{{XL!CELL!SheetName!B5}}`
        * **`{{XL!LAST!cell_ref}}`**: Get the last non-empty value going down from `cell_ref`.
            * Example: `{{XL!LAST!A1}}`
            * Example with sheet: `{{XL!LAST!SheetName!B5}}`
        * **`{{XL!LAST!sheet_name!cell_ref!Title}}`**: Find column by `Title` in the row of `cell_ref`, then get the last non-empty value below `Title`.
            * Example: `{{XL!LAST!Items!A4!Total Project Costs}}`
        * **`{{XL!RANGE!range_ref}}`**: Get values from a range (e.g., `A1:C5` or `NamedRange`). Returns a formatted table.
            * Example: `{{XL!RANGE!A1:C5}}`
            * Example with sheet: `{{XL!RANGE!SheetName!A1:C5}}`
            * Example with named range: `{{XL!RANGE!MyNamedRange}}`
        * **`{{XL!COLUMN!sheet_name!col_refs}}`**: Get specified columns starting from `col_refs` (e.g., "A1,C1,E1"). Returns a table.
            * Example: `{{XL!COLUMN!Items!A4,E4,F4}}`
        * **`{{XL!COLUMN!sheet_name!"Titles"!start_row}}`**: Get columns by `Titles` (e.g., "Revenue,Expense,Profit") found in `start_row`. Returns a table.
            * Example: `{{XL!COLUMN!Items!"Activities,HST,Total Project Costs"!4}}` (Finds titles in row 4)

        ### User Input Keywords (`{{INPUT!...}}`)
        Keywords to create interactive input fields.

        * **`{{INPUT!text!label!default_value}}`**: Single-line text input.
        * **`{{INPUT!area!label!default_value!height}}`**: Multi-line text area (optional height in pixels).
        * **`{{INPUT!date!label!default_date!format}}`**: Date picker (`default_date` can be 'today' or 'YYYY/MM/DD'). `format` is optional (e.g., 'YYYY/MM/DD', 'DD/MM/YYYY').
        * **`{{INPUT!select!label!option1,option2,...}}`**: Dropdown selection.
        * **`{{INPUT!check!label!default_state}}`**: Checkbox (`default_state` is 'True' or 'False').

        ### Template Keywords (`{{TEMPLATE!...}}`)
        Keywords to include content from other files or libraries.

        * **`{{TEMPLATE!filename.docx}}`**: Include entire external template file.
        * **`{{TEMPLATE!filename.docx!section=name}}`**: Include specific section/bookmark.
        * **`{{TEMPLATE!filename.txt!line=5}}`**: Include specific line number from text file.
        * **`{{TEMPLATE!filename.docx!paragraph=3}}`**: Include specific paragraph number.
        * **`{{TEMPLATE!filename.docx!VARS(key1=val1,key2=val2)}}`**: Template with variable substitution (values can be keywords).
        * **`{{TEMPLATE!LIBRARY!template_name!version}}`**: Reference template from a predefined library (optional version).

        ### JSON Data Keywords (`{{JSON!...}}`)
        Keywords to fetch data from JSON files using JSONPath.

        * **`{{JSON!filename.json!json_path}}`**: Access data using JSONPath (e.g., `$.key`, `$.array[0].name`).
            * Example: `{{JSON!config.json!$.settings.theme}}`
            * Example: `{{JSON!data.json!$.users[1].email}}`
        * **`{{JSON!filename.json!json_path!TRANSFORMATION}}`**: Apply optional transformation.
            * `SUM`: Sum numeric values in an array. (`{{JSON!data.json!$.values!SUM}}`)
            * `JOIN(delimiter)`: Join array items with a delimiter. (`{{JSON!data.json!$.names!JOIN(,)}}`)
            * `BOOL(YesText/NoText)`: Transform boolean to custom text. (`{{JSON!config.json!$.enabled!BOOL(Active/Inactive)}}`)
        """
        return help_text 